#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文格式排版引擎
================
读取一份用 Word“内置标题样式”标注好层级的 .docx，
按格式预设（如 GB/T 9704-2012）套用公文版式，输出一份新的 .docx（不改原稿）。

支持：
- 三种字体方案：bundled(随工具内置自动安装) / mac(系统替代) / standard(标准字体名)
- 页边距、各级字体字号、首行缩进、固定行距
- 字体颜色（默认黑色）
- 页码（4号宋体、单页右双页左、一字线包围）
- 可选：自动生成层次序号（一、/（一）/1./（1））
- 元数据清理、校对报告
"""

import json
import os
import re
import shutil
import sys
import tempfile
import uuid
import zipfile
from lxml import etree
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph


def _find_font_dir():
    """查找内置字体目录：按优先级搜索多个可能的位置。

    搜索顺序：
    1. 环境变量 GW_FONT_DIR（用户自定义）
    2. 项目根目录的 fonts/embedded/（开发态）
    3. Skill 自身的 assets/fonts/（打包态）
    4. PyInstaller 的 _MEIPASS/fonts/embedded/（frozen 态）
    """
    candidates = []
    env_dir = os.environ.get("GW_FONT_DIR")
    if env_dir:
        candidates.append(env_dir)
    script_dir = os.path.dirname(os.path.abspath(__file__))
    skill_root = os.path.dirname(script_dir)
    # scripts/ → gongwen-formatter/ → skills/ → .codebuddy/ → 项目根（4 级）
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(script_dir))))
    candidates.append(os.path.join(project_root, "fonts", "embedded"))
    candidates.append(os.path.join(skill_root, "assets", "fonts"))
    if getattr(sys, "frozen", False):
        meipass = getattr(sys, "_MEIPASS", "")
        if meipass:
            candidates.append(os.path.join(meipass, "fonts", "embedded"))
    for d in candidates:
        if d and os.path.isdir(d):
            return d
    return candidates[0] if candidates else ""


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EMBEDDED_FONT_DIR = _find_font_dir()

# 内置字体文件映射：key 必须与 docx 中写入的字体名一致（用于把所用字体嵌入文档）。
_TNR = os.path.join(EMBEDDED_FONT_DIR, "TimesNewRoman", "times.ttf")
FONT_FILE_MAP = {
    "方正小标宋简体": os.path.join(EMBEDDED_FONT_DIR, "方正小标宋简体.ttf"),
    "方正小标宋_GBK": os.path.join(EMBEDDED_FONT_DIR, "方正小标宋_GBK.ttf"),
    "方正大标宋简体": os.path.join(EMBEDDED_FONT_DIR, "方正大标宋简体.ttf"),
    "方正大标宋简繁": os.path.join(EMBEDDED_FONT_DIR, "方正大标宋简繁.ttf"),
    "仿宋_GB2312": os.path.join(EMBEDDED_FONT_DIR, "仿宋_GB2312.ttf"),
    "仿宋": os.path.join(EMBEDDED_FONT_DIR, "仿宋.ttf"),
    "方正仿宋简体": os.path.join(EMBEDDED_FONT_DIR, "方正仿宋简体.ttf"),
    "方正仿宋_GBK": os.path.join(EMBEDDED_FONT_DIR, "方正仿宋_GBK.ttf"),
    "楷体_GB2312": os.path.join(EMBEDDED_FONT_DIR, "楷体_GB2312.ttf"),
    "楷体": os.path.join(EMBEDDED_FONT_DIR, "楷体.ttf"),
    "方正楷体简体": os.path.join(EMBEDDED_FONT_DIR, "方正楷体简体.ttf"),
    "方正楷体_GBK": os.path.join(EMBEDDED_FONT_DIR, "方正楷体_GBK.ttf"),
    "黑体": os.path.join(EMBEDDED_FONT_DIR, "黑体.ttf"),
    "方正黑体简体": os.path.join(EMBEDDED_FONT_DIR, "方正黑体简体.ttf"),
    "方正黑体_GBK": os.path.join(EMBEDDED_FONT_DIR, "方正黑体_GBK.ttf"),
    "宋体": os.path.join(EMBEDDED_FONT_DIR, "宋体.ttc"),
    "Times New Roman": _TNR,
}

# 可供自定义选择的字体名（即上表的键）
SELECTABLE_FONTS = list(FONT_FILE_MAP.keys())


def collect_used_fonts(preset, mode):
    """收集预设在指定字体方案下实际用到的字体名。"""
    names = set()
    for r in preset.get("roles", {}).values():
        if r.get("font") or r.get("bundled_font") or r.get("mac_font"):
            names.add(pick_font(r, mode))
    pn = preset.get("page_number", {})
    if pn.get("enabled"):
        names.add(pick_font(pn, mode))
    lh = preset.get("letterhead", {})
    if lh.get("enabled") and (lh.get("font") or lh.get("bundled_font") or lh.get("mac_font")):
        names.add(pick_font(lh, mode))
    if preset.get("western_font"):
        names.add(preset["western_font"])
    return names


# ---------- 预设加载 ----------
def load_preset(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------- 段落角色识别（依据 Word 内置样式名）----------
_STYLE_MAP = {
    "title": "title", "标题": "title",
    "heading 1": "h1", "标题 1": "h1", "标题1": "h1",
    "heading 2": "h2", "标题 2": "h2", "标题2": "h2",
    "heading 3": "h3", "标题 3": "h3", "标题3": "h3",
    "heading 4": "h4", "标题 4": "h4", "标题4": "h4",
}


def classify_paragraph(paragraph):
    name = (paragraph.style.name or "").strip().lower()
    return _STYLE_MAP.get(name, "body")


# ---------- 按 body 真实顺序遍历段落与表格（用于题注关联）----------
def iter_block_items(parent):
    """按文档中出现的真实顺序产出 Paragraph / Table。"""
    if hasattr(parent, "element"):
        body = parent.element.body
    else:
        body = parent._element
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def is_image_paragraph(paragraph):
    """段落是否承载图片/绘图对象（drawing / pict / object / OLE）。"""
    p = paragraph._p
    return bool(
        p.findall(".//" + qn("w:drawing"))
        or p.findall(".//" + qn("w:object"))
        or p.findall(".//" + qn("w:pict"))
    )


def _match_auto_role(text, roles, order):
    """按 order 中给定的角色顺序，用各角色的 match_regex 匹配文本，命中即返回角色名。

    用于按内容自动识别图题/表题、发文字号、主送机关、落款日期等特殊段落。
    """
    t = (text or "").strip()
    if not t:
        return None
    for key in order:
        cfg = roles.get(key)
        if not cfg:
            continue
        pat = cfg.get("match_regex")
        if pat and re.search(pat, t):
            return key
    return None


# ---------- 字体方案选择 ----------
def pick_font(cfg, mode):
    if mode == "mac" and cfg.get("mac_font"):
        return cfg["mac_font"]
    if mode == "bundled" and cfg.get("bundled_font"):
        return cfg["bundled_font"]
    return cfg["font"]


# ---------- 底层格式工具 ----------
ALIGN_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _set_run_font(run, cn_font, size_pt, bold, color=None, western=None):
    """设置字体：中文(eastAsia)用 cn_font；西文/数字(ascii/hAnsi)用 western(默认同 cn_font)。"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color and color.lower() not in ("auto", "none", ""):
        try:
            run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        except Exception:
            pass
    west = western or cn_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:cs"), west)


def _set_first_line_indent(paragraph, chars, size_pt):
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if chars and chars > 0:
        ind.set(qn("w:firstLineChars"), str(int(chars * 100)))
        ind.set(qn("w:firstLine"), str(int(size_pt * chars * 20)))
    else:
        for a in ("w:firstLineChars", "w:firstLine"):
            ind.attrib.pop(qn(a), None)


def _set_line_spacing(paragraph, line_pt):
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(line_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def _fix_smart_quotes(paragraph):
    """把段落中的英文直双引号 " 按出现顺序配对转为中文弯引号 “ ”。

    仅处理中文公文常见的双引号；单引号因易与英文撇号混淆，暂不处理。
    逐 run 处理以保留各 run 的字体/加粗等格式。
    """
    open_quote = True
    for run in paragraph.runs:
        if '"' not in run.text:
            continue
        out = []
        for ch in run.text:
            if ch == '"':
                out.append("\u201c" if open_quote else "\u201d")
                open_quote = not open_quote
            else:
                out.append(ch)
        run.text = "".join(out)


def _apply_role(paragraph, role_cfg, line_pt, mode, western=None, preserve_bold=False):
    font = pick_font(role_cfg, mode)
    size = role_cfg["size_pt"]
    bold = role_cfg.get("bold", False)
    color = role_cfg.get("color", "000000")
    paragraph.alignment = ALIGN_MAP.get(role_cfg.get("align", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY)
    _set_line_spacing(paragraph, line_pt)
    _set_first_line_indent(paragraph, role_cfg.get("first_line_indent_chars", 0), size)
    pf = paragraph.paragraph_format
    if "space_before_pt" in role_cfg:
        pf.space_before = Pt(role_cfg["space_before_pt"])
    if "space_after_pt" in role_cfg:
        pf.space_after = Pt(role_cfg["space_after_pt"])
    for run in paragraph.runs:
        # 仅正文保留原稿加粗（强调词）；标题/各级标题用预设的 bold 值
        run_bold = (bold or (run.font.bold or False)) if preserve_bold else bold
        _set_run_font(run, font, size, run_bold, color, western)


def _apply_figure(paragraph, fig_cfg):
    """图片段落：居中、去首行缩进、单倍行距（切勿用固定行距，否则裁切图片）。"""
    paragraph.alignment = ALIGN_MAP.get(fig_cfg.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
    _set_first_line_indent(paragraph, 0, 16)
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if "space_before_pt" in fig_cfg:
        pf.space_before = Pt(fig_cfg["space_before_pt"])
    if "space_after_pt" in fig_cfg:
        pf.space_after = Pt(fig_cfg["space_after_pt"])


# ---------- 三线表 ----------
def _border_el(tag, pt, color="000000"):
    e = OxmlElement("w:" + tag)
    if pt and pt > 0:
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(int(round(pt * 8))))  # OOXML 边框单位 = 1/8 pt
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
    else:
        e.set(qn("w:val"), "nil")
    return e


def _set_table_borders(table, top_pt, bottom_pt, inner):
    """设置表格整体边框：顶/底为实线，左右与竖线去除，内横线按 inner。"""
    tblPr = table._tbl.tblPr
    for b in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(b)
    borders = OxmlElement("w:tblBorders")
    inner_pt = top_pt if inner else 0
    for tag, pt in (("top", top_pt), ("left", 0), ("bottom", bottom_pt),
                    ("right", 0), ("insideH", inner_pt), ("insideV", 0)):
        borders.append(_border_el(tag, pt))
    tblPr.append(borders)


def _clear_cell_borders(table):
    """清除单元格级自带边框（如 Table Grid 样式遗留），让表级边框生效。"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for b in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(b)


def _set_header_bottom(table, pt):
    """给表头行（首行）加一条栏目线（底边框）。"""
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        for b in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(b)
        tb = OxmlElement("w:tcBorders")
        tb.append(_border_el("bottom", pt))
        tcPr.append(tb)


def _apply_three_line_table(table, ts):
    """把普通表格规范化为三线表。返回是否处理。"""
    if not ts or not ts.get("three_line"):
        return False
    try:
        table.style = None
    except Exception:
        pass
    _clear_cell_borders(table)
    outer = ts.get("outer_line_pt", 1.5)
    _set_table_borders(table, outer, outer, ts.get("inner_lines", False))
    _set_header_bottom(table, ts.get("header_line_pt", 0.75))
    if table.rows and ts.get("header_bold", True):
        hdr_align = ALIGN_MAP.get(ts.get("header_align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                p.alignment = hdr_align
                for run in p.runs:
                    run.font.bold = True
    return True


# ---------- 红头（发文机关标志 + 红色分隔线）----------
def _set_paragraph_border(paragraph, side, color, width_pt):
    """给段落某一侧加边框，用来绘制公文红头/页脚的红色分隔线（红线）。

    side: 'bottom' / 'top'。width_pt 转换为 OOXML 的 1/8 pt 单位。
    """
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    for old in pbdr.findall(qn("w:" + side)):
        pbdr.remove(old)
    el = OxmlElement("w:" + side)
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(max(2, int(round(width_pt * 8)))))  # 1/8 pt
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    pbdr.append(el)


def _insert_letterhead(document, lh, org_name, mode, western):
    """在文档最顶部插入红头：居中的红色发文机关名 + 下方红色分隔线。

    lh: letterhead 配置；org_name: 机关名（界面可覆盖，空则用预设）。
    """
    name = (org_name or lh.get("org_name") or "").strip()
    if not name:
        return False
    font = pick_font(lh, mode)
    size = lh.get("size_pt", 22)
    bold = lh.get("bold", False)
    color = lh.get("color", "FF0000")
    align = lh.get("align", "center")

    # 在首段之前插入两段：机关名段 + 红线段（顺序：机关名 → 红线 → 原内容）
    if document.paragraphs:
        anchor = document.paragraphs[0]
        line_par = anchor.insert_paragraph_before()
        org_par = line_par.insert_paragraph_before()
    else:
        org_par = document.add_paragraph()
        line_par = document.add_paragraph()

    # 机关名段
    org_par.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    pf = org_par.paragraph_format
    pf.space_before = Pt(lh.get("space_before_pt", 6))
    pf.space_after = Pt(lh.get("space_after_pt", 2))
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = org_par.add_run(name)
    _set_run_font(run, font, size, bold, color, western)

    # 红线段（空段落 + 红色底边框）
    line_cfg = lh.get("line", {})
    if line_cfg.get("enabled", True):
        line_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lpf = line_par.paragraph_format
        lpf.space_before = Pt(line_cfg.get("space_before_pt", 0))
        lpf.space_after = Pt(line_cfg.get("space_after_pt", 10))
        lpf.line_spacing = 1.0
        lpf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # 让空段尽量矮：设一个很小的字号占位 run
        holder = line_par.add_run("")
        holder.font.size = Pt(1)
        _set_paragraph_border(
            line_par, "bottom", line_cfg.get("color", color), line_cfg.get("width_pt", 3))
    else:
        # 不画线时移除多余空段
        line_par._p.getparent().remove(line_par._p)
    return True


def _add_footer_line(document, cfg):
    """在页脚处绘制一条红色横线（公文版心底部红线），用段落上边框实现。"""
    color = cfg.get("color", "FF0000")
    width_pt = cfg.get("width_pt", 3)
    dist_mm = cfg.get("distance_mm")
    for section in document.sections:
        if dist_mm:
            section.footer_distance = Mm(dist_mm)
        footer = section.footer
        footer.is_linked_to_previous = False
        par = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = par.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if not par.runs:
            r = par.add_run("")
            r.font.size = Pt(1)
        _set_paragraph_border(par, "top", color, width_pt)


def _set_page(section, page):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(page["width_mm"])
    section.page_height = Mm(page["height_mm"])
    section.top_margin = Mm(page["margin_top_mm"])
    section.bottom_margin = Mm(page["margin_bottom_mm"])
    section.left_margin = Mm(page["margin_left_mm"])
    section.right_margin = Mm(page["margin_right_mm"])


def _clean_metadata(document):
    cp = document.core_properties
    for attr in ("author", "last_modified_by", "comments", "category",
                 "keywords", "subject", "title", "company"):
        try:
            setattr(cp, attr, "")
        except Exception:
            pass


# ---------- 字体嵌入（解决目标电脑未安装/字体缓存未刷新导致不显示的问题） ----------
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
FONT_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"
FONT_CT = "application/vnd.openxmlformats-officedocument.obfuscatedFont"
XML_PARSER = etree.XMLParser(remove_blank_text=False)


def _xml(data):
    return etree.fromstring(data, parser=XML_PARSER)


def _xml_bytes(root):
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def _obfuscate_font(data, font_key):
    """按 OOXML 规则混淆字体：前32字节与 GUID(bytes_le) 循环异或。"""
    b = bytearray(data)
    key = font_key.bytes_le
    for i in range(min(32, len(b))):
        b[i] ^= key[i % 16]
    return bytes(b)


def _next_rid(rels_root):
    used = set()
    for rel in rels_root.findall(f"{{{REL_NS}}}Relationship"):
        rid = rel.get("Id", "")
        if rid.startswith("rId") and rid[3:].isdigit():
            used.add(int(rid[3:]))
    n = 1
    while n in used:
        n += 1
    return f"rId{n}"


def _ensure_font_node(fonts_root, font_name):
    for node in fonts_root.findall(f"{{{W_NS}}}font"):
        if node.get(f"{{{W_NS}}}name") == font_name:
            return node
    node = etree.SubElement(fonts_root, f"{{{W_NS}}}font")
    node.set(f"{{{W_NS}}}name", font_name)
    etree.SubElement(node, f"{{{W_NS}}}charset").set(f"{{{W_NS}}}val", "86")
    etree.SubElement(node, f"{{{W_NS}}}family").set(f"{{{W_NS}}}val", "roman")
    return node


def _embed_fonts_in_docx(docx_path, font_map):
    """把 font_map 中存在的字体嵌入到 docx。返回成功嵌入的字体名列表。"""
    existing = {name: path for name, path in font_map.items() if os.path.exists(path)}
    if not existing:
        return []

    with zipfile.ZipFile(docx_path, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}

    # Content Types
    ct_root = _xml(parts["[Content_Types].xml"])
    if not any(x.get("Extension") == "odttf" for x in ct_root.findall(f"{{{CT_NS}}}Default")):
        d = etree.SubElement(ct_root, f"{{{CT_NS}}}Default")
        d.set("Extension", "odttf")
        d.set("ContentType", FONT_CT)
    parts["[Content_Types].xml"] = _xml_bytes(ct_root)

    # fontTable.xml
    font_table_path = "word/fontTable.xml"
    fonts_root = _xml(parts.get(font_table_path, b'<w:fonts xmlns:w="%s" xmlns:r="%s"/>' % (W_NS.encode(), R_NS.encode())))

    # fontTable relationships
    rels_path = "word/_rels/fontTable.xml.rels"
    if rels_path in parts:
        rels_root = _xml(parts[rels_path])
    else:
        rels_root = etree.Element(f"{{{REL_NS}}}Relationships")

    embedded = []
    for idx, (font_name, font_path) in enumerate(existing.items(), start=1):
        font_key = uuid.uuid4()
        rid = _next_rid(rels_root)
        target = f"fonts/font{idx}.odttf"
        with open(font_path, "rb") as f:
            parts[f"word/{target}"] = _obfuscate_font(f.read(), font_key)

        rel = etree.SubElement(rels_root, f"{{{REL_NS}}}Relationship")
        rel.set("Id", rid)
        rel.set("Type", FONT_REL_TYPE)
        rel.set("Target", target)

        font_node = _ensure_font_node(fonts_root, font_name)
        for child_name in ("embedRegular", "embedBold", "embedItalic", "embedBoldItalic"):
            for old in font_node.findall(f"{{{W_NS}}}{child_name}"):
                font_node.remove(old)
        emb = etree.SubElement(font_node, f"{{{W_NS}}}embedRegular")
        emb.set(f"{{{R_NS}}}id", rid)
        emb.set(f"{{{W_NS}}}fontKey", "{" + str(font_key).upper() + "}")
        embedded.append(font_name)

    parts[font_table_path] = _xml_bytes(fonts_root)
    parts[rels_path] = _xml_bytes(rels_root)

    # settings: 明确启用嵌入字体
    settings_path = "word/settings.xml"
    if settings_path in parts:
        settings = _xml(parts[settings_path])
        if settings.find(f"{{{W_NS}}}embedTrueTypeFonts") is None:
            settings.append(etree.Element(f"{{{W_NS}}}embedTrueTypeFonts"))
        parts[settings_path] = _xml_bytes(settings)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    try:
        with zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in parts.items():
                zout.writestr(name, data)
        shutil.move(tmp.name, docx_path)
    finally:
        if os.path.exists(tmp.name):
            os.remove(tmp.name)
    return embedded


# ---------- 页码 ----------
def _enable_even_odd(document):
    settings = document.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def _fill_pagenum_paragraph(paragraph, align, fmt, font, size, color):
    # 清空已有 runs
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    paragraph.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    prefix, suffix = (fmt.split("{n}") + [""])[:2] if "{n}" in fmt else ("", "")
    if prefix:
        run = paragraph.add_run(prefix)
        _set_run_font(run, font, size, False, color)
    # PAGE 域
    run = paragraph.add_run()
    _set_run_font(run, font, size, False, color)
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)
    if suffix:
        run = paragraph.add_run(suffix)
        _set_run_font(run, font, size, False, color)


def _add_page_numbers(document, pn_cfg, mode):
    if not pn_cfg or not pn_cfg.get("enabled"):
        return
    _enable_even_odd(document)
    font = pick_font(pn_cfg, mode)
    size = pn_cfg.get("size_pt", 14)
    fmt = pn_cfg.get("format", "— {n} —")
    color = pn_cfg.get("color", "000000")
    for section in document.sections:
        section.footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False
        odd = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        _fill_pagenum_paragraph(odd, pn_cfg.get("odd_align", "right"), fmt, font, size, color)
        ev = section.even_page_footer
        evp = ev.paragraphs[0] if ev.paragraphs else ev.add_paragraph()
        _fill_pagenum_paragraph(evp, pn_cfg.get("even_align", "left"), fmt, font, size, color)


# ---------- 可选：自动层次序号 ----------
_CN_NUM = "零一二三四五六七八九"


def _int_to_cn(n):
    if n <= 10:
        return "十" if n == 10 else _CN_NUM[n]
    if n < 20:
        return "十" + _CN_NUM[n - 10]
    if n < 100:
        return _CN_NUM[n // 10] + "十" + (_CN_NUM[n % 10] if n % 10 else "")
    return str(n)


# 去除作者手输的各类前导序号
_STRIP_PATTERNS = [
    r"^\s*[（(][一二三四五六七八九十百零\d]+[)）]\s*",
    r"^\s*[一二三四五六七八九十百零]+\s*[、.．]\s*",
    r"^\s*\d+\s*[、.．]\s*",
]


def _strip_leading_number(text):
    for pat in _STRIP_PATTERNS:
        text = re.sub(pat, "", text, count=1)
    return text


def _make_number(style, idx):
    if style == "cn_dun":
        return f"{_int_to_cn(idx)}、"
    if style == "cn_paren":
        return f"（{_int_to_cn(idx)}）"
    if style == "digit_dot":
        return f"{idx}."
    if style == "digit_paren":
        return f"（{idx}）"
    return ""


def _apply_auto_numbering(doc, roles):
    counters = {"h1": 0, "h2": 0, "h3": 0, "h4": 0}
    order = ["h1", "h2", "h3", "h4"]
    for p in doc.paragraphs:
        role = classify_paragraph(p)
        if role not in counters:
            continue
        style = roles[role].get("auto_number")
        if not style:
            continue
        counters[role] += 1
        # 重置更深层级计数
        for deeper in order[order.index(role) + 1:]:
            counters[deeper] = 0
        if not p.runs:
            continue
        num = _make_number(style, counters[role])
        # 在第一个 run 前加序号，并去掉原有手输序号
        first = p.runs[0]
        first.text = num + _strip_leading_number(first.text)


# ---------- 主流程 ----------
def format_document(input_path, preset, output_path=None,
                    font_mode=None, clean_metadata=True, auto_number=False,
                    letterhead_org=None):
    doc = Document(input_path)
    page = preset["page"]
    line_pt = preset.get("line_spacing_pt", 28.8)
    roles = preset["roles"]
    mode = font_mode or preset.get("font_mode_default", "bundled")
    western = preset.get("western_font")
    # 按内容自动识别的角色顺序（图题/表题 + 发文字号/主送/落款等）
    auto_match = preset.get("auto_match_roles", ["fig_caption", "tbl_caption"])

    for section in doc.sections:
        _set_page(section, page)

    table_style = preset.get("table_style", {})
    counts = {"title": 0, "h1": 0, "h2": 0, "h3": 0, "h4": 0, "body": 0,
              "figure": 0, "fig_caption": 0, "tbl_caption": 0}
    table_count = 0
    three_line_count = 0

    # 红头模式下，第一个内容段落自动作为大标题（不论 Word 样式）
    lh_cfg = preset.get("letterhead", {})
    letterhead_enabled = lh_cfg.get("enabled", False)

    # ===== 第一遍：处理表格/图片，收集文字段落 =====
    text_paras = []
    for block in iter_block_items(doc):
        if isinstance(block, Table):
            table_count += 1
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _apply_role(p, roles["table"], line_pt, mode, western, preserve_bold=True)
                        _fix_smart_quotes(p)
            if _apply_three_line_table(block, table_style):
                three_line_count += 1
            continue
        if is_image_paragraph(block):
            if "figure" in roles:
                _apply_figure(block, roles["figure"])
                counts["figure"] += 1
            continue
        if block.text.strip() or block.runs:
            text_paras.append(block)

    n = len(text_paras)
    assigned = [None] * n

    # ===== 第二遍：用「位置 + 内容」双重启发式确定角色 =====
    # (a) 大标题：红头模式下首段即标题；否则按 Word 样式
    title_idx = None
    if letterhead_enabled and n > 0:
        assigned[0] = "title"
        title_idx = 0
    else:
        for i, p in enumerate(text_paras):
            if classify_paragraph(p) == "title":
                assigned[i] = "title"
                title_idx = i
                break

    # (b) 主送机关/称呼语：标题后第一个段落，以冒号结尾（位置启发式，兼容“尊敬的各位领导：”）
    if "recipient" in roles and title_idx is not None:
        j = title_idx + 1
        if j < n and assigned[j] is None:
            t = (text_paras[j].text or "").strip()
            if t and len(t) <= 40 and re.search(r"[：:]\s*$", t):
                assigned[j] = "recipient"

    # (c) 落款：从文末向上找日期，其紧邻上一段作为署名（位置启发式）
    date_re = roles.get("signoff", {}).get("match_regex")
    if "signoff" in roles and date_re:
        for i in range(n - 1, -1, -1):
            if assigned[i] is not None:
                continue
            t = (text_paras[i].text or "").strip()
            if t and re.search(date_re, t):
                assigned[i] = "signoff"
                k = i - 1
                while k >= 0 and not (text_paras[k].text or "").strip():
                    k -= 1
                if k > (title_idx or 0) and assigned[k] is None and "signoff_org" in roles:
                    tk = (text_paras[k].text or "").strip()
                    if tk and len(tk) <= 40 and not re.search(r"[。！？!?]\s*$", tk):
                        assigned[k] = "signoff_org"
                break

    # (d) 其余段落：Word 样式 + 内容匹配（仅图题/表题/发文字号，位置型角色不参与内容匹配）
    content_match = [r for r in auto_match if r in ("doc_number", "fig_caption", "tbl_caption")]
    for i, p in enumerate(text_paras):
        if assigned[i] is not None:
            continue
        role = classify_paragraph(p)
        if role == "body":
            matched = _match_auto_role(p.text, roles, content_match)
            if matched:
                role = matched
        assigned[i] = role

    # ===== 第三遍：套用格式 + 智能引号 =====
    for i, p in enumerate(text_paras):
        role = assigned[i] if assigned[i] in roles else "body"
        # 落款署名与日期之间插入空行（段落标记）
        if role == "signoff" and i > 0 and assigned[i - 1] == "signoff_org":
            blank = p.insert_paragraph_before()
            blank.paragraph_format.line_spacing = Pt(line_pt)
            blank.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        _apply_role(p, roles[role], line_pt, mode, western, preserve_bold=(role == "body"))
        _fix_smart_quotes(p)
        counts[role] = counts.get(role, 0) + 1

    if auto_number:
        _apply_auto_numbering(doc, roles)

    # 红头（发文机关标志 + 红线）：在所有正文格式化之后插入到文档最顶部
    letterhead_done = False
    lh_cfg = preset.get("letterhead", {})
    if lh_cfg.get("enabled"):
        letterhead_done = _insert_letterhead(doc, lh_cfg, letterhead_org, mode, western)

    # 页脚红线（版心底部红线）
    footer_line_done = False
    fl_cfg = preset.get("footer_line", {})
    if fl_cfg.get("enabled"):
        _add_footer_line(doc, fl_cfg)
        footer_line_done = True

    _add_page_numbers(doc, preset.get("page_number"), mode)

    if clean_metadata:
        _clean_metadata(doc)

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = base + "-公文版.docx"
    doc.save(output_path)

    embedded_fonts = []
    if mode == "bundled":
        used = collect_used_fonts(preset, mode)
        fmap = {n: FONT_FILE_MAP[n] for n in used if n in FONT_FILE_MAP}
        embedded_fonts = _embed_fonts_in_docx(output_path, fmap)

    mode_label = {"bundled": "内置公文字体（已嵌入文档）", "mac": "Mac系统替代字体", "standard": "标准字体名"}.get(mode, mode)
    warnings = []
    if letterhead_done:
        shown = (letterhead_org or lh_cfg.get("org_name") or "").strip()
        warnings.append(f"已在文首插入红头：红色机关名“{shown}”+ 红色分隔线。如需修改机关名，可在排版时指定机关名称。")
    if footer_line_done:
        warnings.append("已在版心底部添加页脚红线。")
    sp_msgs = []
    for key, cn in (("doc_number", "发文字号"), ("recipient", "主送机关"),
                    ("signoff_org", "落款署名"), ("signoff", "落款日期")):
        if counts.get(key):
            sp_msgs.append(f"{cn} {counts[key]} 处")
    if sp_msgs:
        warnings.append("已按内容自动识别并排版：" + "、".join(sp_msgs) + "（发文字号/落款右对齐、主送机关顶格、落款署名与日期间空一行）。请人工核对。")
    if counts["title"] == 0:
        warnings.append("未检测到“标题”样式的文章大标题（可在 Word 中把大标题设为“标题”样式）。")
    if counts["title"] > 1:
        warnings.append(f"检测到 {counts['title']} 个大标题，公文通常只应有 1 个，请核对。")
    if counts["h1"] + counts["h2"] + counts["h3"] + counts["h4"] == 0:
        warnings.append("未检测到任何层次标题，全部按正文处理；如有层级请用 Word“标题1/2/3/4”标注。")
    if table_count > 0:
        if three_line_count > 0:
            warnings.append(f"检测到 {table_count} 个表格，已统一表内字体并转为三线表（{three_line_count} 个，表头加粗居中），请人工核对表头与对齐。")
        else:
            warnings.append(f"检测到 {table_count} 个表格，已统一表内字体（未启用三线表）。")
    if counts["figure"] > 0:
        warnings.append(f"检测到 {counts['figure']} 张图片，已居中并去除首行缩进。")
    if counts["fig_caption"] + counts["tbl_caption"] > 0:
        warnings.append(f"已规范化图题 {counts['fig_caption']} 处、表题 {counts['tbl_caption']} 处（黑体居中）；如有漏判请确认题注以“图/表+编号”开头。")
    if preset.get("page_number", {}).get("enabled"):
        warnings.append("已添加页码（单页右、双页左）。在 Word 中查看页码可能需先更新域：全选→F9（Mac：fn+F9）。")
    if auto_number:
        warnings.append("已自动生成层次序号（一、/（一）/1./（1）），并替换了原手输序号，请核对是否正确。")
    if mode == "bundled":
        if embedded_fonts:
            warnings.append("已将内置公文字体嵌入到输出文档，其他电脑即使未安装字体也应能正常显示。")
        else:
            warnings.append("未找到可嵌入的内置字体文件，请检查 fonts/embedded 目录。")

    report = {
        "input": input_path, "output": output_path, "font_mode": mode_label,
        "counts": counts, "tables": table_count, "metadata_cleaned": clean_metadata,
        "auto_number": auto_number, "embedded_fonts": embedded_fonts, "warnings": warnings,
    }
    return output_path, report


# ---------- 预设详情描述（用于界面提示栏） ----------
_PT_TO_HAO = {42: "初号", 36: "小初", 26: "一号", 24: "小一", 22: "二号",
              18: "小二", 16: "三号", 15: "小三", 14: "四号", 12: "小四",
              10.5: "五号", 9: "小五"}


def _size_label(pt):
    hao = _PT_TO_HAO.get(pt)
    return f"{hao}（{pt:g}pt）" if hao else f"{pt:g}pt"


def _align_label(a):
    return {"center": "居中", "left": "左对齐", "right": "右对齐", "justify": "两端对齐"}.get(a, a)


def describe_preset(preset, mode="bundled"):
    """把预设的关键排版参数整理成可读文本，显示在界面提示栏。"""
    pg = preset.get("page", {})
    roles = preset.get("roles", {})
    lines = [f"【当前格式】{preset.get('preset_name', '')}"]
    lines.append("")
    lines.append("· 纸张：%s，页边距 上%s/下%s/左%s/右%s mm" % (
        pg.get("size", "A4"), pg.get("margin_top_mm"), pg.get("margin_bottom_mm"),
        pg.get("margin_left_mm"), pg.get("margin_right_mm")))
    lines.append("· 行距：固定值 %g 磅" % preset.get("line_spacing_pt", 28.8))
    if preset.get("western_font"):
        lines.append("· 西文/数字：%s" % preset["western_font"])

    lh = preset.get("letterhead", {})
    if lh.get("enabled"):
        seg = "· 红头：%s %s 红色 %s" % (
            pick_font(lh, mode), _size_label(lh.get("size_pt", 22)),
            _align_label(lh.get("align", "center")))
        if lh.get("line", {}).get("enabled", True):
            seg += " + 红色分隔线"
        lines.append(seg)
        if lh.get("org_name"):
            lines.append("  （默认机关名：%s，可在主界面修改）" % lh["org_name"])

    order = [("title", "大标题"), ("h1", "一级标题"), ("h2", "二级标题"),
             ("h3", "三级标题"), ("h4", "四级标题"), ("body", "正文"), ("table", "表格文字")]
    for key, cn in order:
        r = roles.get(key)
        if not r:
            continue
        font = pick_font(r, mode)
        seg = "· %s：%s %s" % (cn, font, _size_label(r.get("size_pt", 16)))
        if r.get("bold"):
            seg += " 加粗"
        seg += " " + _align_label(r.get("align", "justify"))
        if r.get("first_line_indent_chars"):
            seg += " 首行缩进%d字" % r["first_line_indent_chars"]
        if r.get("numbering"):
            seg += "（序号 %s）" % r["numbering"]
        lines.append(seg)

    pn = preset.get("page_number", {})
    if pn.get("enabled"):
        lines.append("· 页码：%s %s，单页%s/双页%s，样式“%s”" % (
            pick_font(pn, mode), _size_label(pn.get("size_pt", 14)),
            _align_label(pn.get("odd_align", "right")), _align_label(pn.get("even_align", "left")),
            pn.get("format", "— {n} —").replace("{n}", "页码")))
    return "\n".join(lines)


def report_to_text(report):
    c = report["counts"]
    lines = [
        "==== 公文排版完成 ====",
        f"输入文件：{os.path.basename(report['input'])}",
        f"输出文件：{os.path.basename(report['output'])}",
        f"字体方案：{report['font_mode']}",
        f"嵌入字体：{('、'.join(report.get('embedded_fonts', [])) or '无')}",
        f"元数据清理：{'已清除作者/单位等信息' if report['metadata_cleaned'] else '未清理'}",
        f"自动层次序号：{'已开启' if report['auto_number'] else '关闭'}",
        "",
        "段落统计：",
        f"  文章大标题：{c['title']}",
        f"  一级标题：{c['h1']}    二级标题：{c['h2']}",
        f"  三级标题：{c['h3']}    四级标题：{c['h4']}",
        f"  正文段落：{c['body']}",
        f"  图片：{c.get('figure', 0)}    图题：{c.get('fig_caption', 0)}    表题：{c.get('tbl_caption', 0)}",
        f"  表格数量：{report['tables']}",
    ]
    if report["warnings"]:
        lines.append("")
        lines.append("提醒（请人工核对）：")
        for w in report["warnings"]:
            lines.append(f"  · {w}")
    else:
        lines.append("\n未发现明显问题。")
    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python format_engine.py 输入文件.docx [预设.json] [--mode bundled|mac|standard] [--number]")
        sys.exit(1)
    inp = sys.argv[1]
    preset_path = next((a for a in sys.argv[2:] if a.endswith(".json")),
                       os.path.join(os.path.dirname(__file__), "..", "assets", "presets", "redhead_notice.json"))
    mode = None
    if "--mode" in sys.argv:
        mode = sys.argv[sys.argv.index("--mode") + 1]
    auto_num = "--number" in sys.argv
    out, rep = format_document(inp, load_preset(preset_path), font_mode=mode, auto_number=auto_num)
    print(report_to_text(rep))
